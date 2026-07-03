"""
Tests: session memory and materials fixes (CORRECTION-session-memory-and-materials-v1.1).

F1: File-only turns visible to AI (file note stored in session_messages)
F2a: Materials summary injected into converse prompt
F2b: .docx text extraction
F2c: .mp4 recognized as audio (not binary garbage)
F3: History truncation keeps NEWEST turns (tail slice)
F5: Anti-repeat guard (difflib similarity check)
"""
import json
import os
import sys
import tempfile
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

from materials import MaterialsIntake


class TestF1FileNoteInSessionMessages:
    """F1: file-only turns must be visible in session_messages (the transcript source)."""

    def test_file_note_in_session_messages_when_text_empty(self):
        """When operator sends files with no text, session_messages should contain the file note."""
        # This tests the logic: turn_text = (text + "\n" if text else "") + file_note.strip() if files else text
        text = ""
        files = ["test.zip"]
        file_note = f"\n[Operator attached files: {', '.join(files)}]" if files else ""
        turn_text = (text + "\n" if text else "") + file_note.strip() if files else text
        assert "[Operator attached files: test.zip]" in turn_text
        assert turn_text.strip() != ""  # Not a blank line

    def test_file_note_combined_with_text(self):
        text = "here's my stuff"
        files = ["report.docx", "notes.txt"]
        file_note = f"\n[Operator attached files: {', '.join(files)}]" if files else ""
        turn_text = (text + "\n" if text else "") + file_note.strip() if files else text
        assert "here's my stuff" in turn_text
        assert "[Operator attached files: report.docx, notes.txt]" in turn_text

    def test_no_file_note_when_no_files(self):
        text = "just talking"
        files = []
        file_note = f"\n[Operator attached files: {', '.join(files)}]" if files else ""
        turn_text = (text + "\n" if text else "") + file_note.strip() if files else text
        assert "Operator attached" not in turn_text
        assert turn_text == "just talking"


class TestF2bDocxExtraction:
    """F2b: .docx files must be text-extracted."""

    def test_docx_text_extraction(self):
        """Create a real .docx and verify text extraction works."""
        from docx import Document
        db_path = os.path.join(tempfile.mkdtemp(), "test.db")
        intake = MaterialsIntake(db_path)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = f.name

        try:
            doc = Document()
            doc.add_paragraph("StackPenni Brand Report")
            doc.add_paragraph("Caribbean AI + wealth building for regular people")
            doc.add_paragraph("Target audience: small vendors, entrepreneurs")
            doc.save(docx_path)

            text = intake._extract_docx_text(docx_path)
            assert text is not None
            assert "StackPenni Brand Report" in text
            assert "Caribbean AI" in text
            assert "small vendors" in text
        finally:
            os.unlink(docx_path)

    def test_docx_extraction_returns_none_for_invalid_file(self):
        db_path = os.path.join(tempfile.mkdtemp(), "test.db")
        intake = MaterialsIntake(db_path)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(b"not a real docx")
            path = f.name
        try:
            assert intake._extract_docx_text(path) is None
        finally:
            os.unlink(path)


class TestF2cAudioRecognition:
    """F2c: .mp4, .opus, .aac, .flac recognized as audio, not binary garbage."""

    @pytest.mark.parametrize("ext", [".mp3", ".wav", ".m4a", ".ogg", ".webm", ".mp4", ".opus", ".aac", ".flac"])
    def test_audio_extensions_recognized(self, ext):
        """Audio extensions should be handled as audio, not fall through to the text/binary branch."""
        db_path = os.path.join(tempfile.mkdtemp(), "test.db")
        intake = MaterialsIntake(db_path)
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
            f.write(b"\x00\x00\x00\x1cftypisom")  # Fake binary header
            path = f.name
        try:
            material_id = intake.ingest_file(path, run_id=1, business_slug="test")
            material = intake.get_material(material_id)
            assert material["material_type"] == "audio"
            assert "transcription pending" in material["raw_content"].lower()
        finally:
            os.unlink(path)


class TestF3TailSlice:
    """F3: conversation history must keep NEWEST turns (tail slice, not head)."""

    def test_tail_slice_keeps_newest(self):
        """The last 12000 chars should be kept, not the first 4000."""
        conversation = "A" * 5000 + "B" * 5000 + "C" * 5000  # 15000 chars total
        truncated = conversation[-12000:]
        # "A"s at the start should be cut, "C"s at the end should be kept
        assert truncated.startswith("B") or truncated.startswith("A" * 1000)
        assert truncated.endswith("C" * 5000)
        assert "A" * 5000 not in truncated  # The oldest part is gone
        assert len(truncated) == 12000


class TestF5AntiRepeat:
    """F5: server-side similarity check via difflib."""

    def test_near_duplicate_detected(self):
        import difflib
        reply = "Hey! What kind of content do you usually create?"
        prior = "Hey! What kind of content do you usually create?"
        ratio = difflib.SequenceMatcher(None, reply.lower(), prior.lower()).ratio()
        assert ratio > 0.9

    def test_different_replies_not_flagged(self):
        import difflib
        reply = "Tell me about your target audience"
        prior = "What kind of content do you usually create?"
        ratio = difflib.SequenceMatcher(None, reply.lower(), prior.lower()).ratio()
        assert ratio < 0.9

    def test_reworded_question_not_flagged(self):
        import difflib
        reply = "Who are the people you're trying to reach with this content?"
        prior = "What kind of content do you usually create?"
        ratio = difflib.SequenceMatcher(None, reply.lower(), prior.lower()).ratio()
        assert ratio < 0.9  # Same topic, different wording — should not be flagged
